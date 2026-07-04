import os
import io
import fitz  # PyMuPDF (used to render PDF pages to images for OCR)
import pdfplumber
import pytesseract
import openpyxl
from PIL import Image
from docx import Document

# Set up Tesseract binary path for default Windows installation
tesseract_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
if os.path.exists(tesseract_path):
    pytesseract.pytesseract.tesseract_cmd = tesseract_path

def extract_pdf(path):
    """Extracts text from a PDF file using pdfplumber. 
    Falls back to pytesseract OCR if the file is a scanned PDF (no text found).
    """
    text = ""
    
    # 1. Try standard text extraction
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
                
    # 2. If no text was found (scanned PDF), run OCR
    if not text.strip():
        try:
            doc = fitz.open(path)
            for page in doc:
                # Render PDF page to image bytes (150 DPI is recommended for OCR)
                pix = page.get_pixmap(dpi=150)
                img_data = pix.tobytes("png")
                # Load image in PIL
                img = Image.open(io.BytesIO(img_data))
                # Run OCR text extraction
                page_text = pytesseract.image_to_string(img)
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            print(f"[Warning] OCR extraction failed for {path} (check Tesseract installation): {e}")
            
    return text

def extract_word(path):
    """Extracts text from a Word document (.docx) paragraph-by-paragraph and table-by-table."""
    text = ""
    doc = Document(path)
    
    # Extract paragraph texts
    for p in doc.paragraphs:
        if p.text.strip():
            text += p.text + "\n"
            
    # Extract table texts
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            text += " | ".join(row_text) + "\n"
            
    return text

def extract_excel(path):
    """Extracts text sheet-by-sheet from an Excel spreadsheet (.xlsx, .xls) using openpyxl."""
    text = ""
    wb = openpyxl.load_workbook(path, data_only=True)
    
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        text += f"## Sheet: {sheet_name}\n"
        
        # Iterate row by row and join cell values
        for row in sheet.iter_rows(values_only=True):
            if any(cell_val is not None for cell_val in row):
                row_str = " | ".join(str(val).strip() if val is not None else "" for val in row)
                text += row_str + "\n"
        text += "\n"
        
    return text

if __name__ == "__main__":
    import sys
    # Configure stdout to use UTF-8 encoding on Windows to prevent UnicodeEncodeError in console print statements
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except AttributeError:
        pass

    data_dir = "data"
    if not os.path.exists(data_dir):
        print(f"Error: Data directory '{data_dir}' does not exist.")
    else:
        pdf_exts = {".pdf"}
        word_exts = {".docx", ".doc"}
        excel_exts = {".xlsx", ".xls"}
        
        # Recursively find and process all files in data directory
        for root, _, files in os.walk(data_dir):
            for file in files:
                file_path = os.path.join(root, file)
                _, ext = os.path.splitext(file.lower())
                
                print(f"\n==================================================")
                print(f"Extracting data from: {file_path}")
                print(f"==================================================")
                
                try:
                    if ext in pdf_exts:
                        text = extract_pdf(file_path)
                    elif ext in word_exts:
                        text = extract_word(file_path)
                    elif ext in excel_exts:
                        text = extract_excel(file_path)
                    else:
                        print(f"Unsupported format: {ext}")
                        continue
                        
                    print(f"Successfully extracted {len(text)} characters.")
                    print("--- Preview (First 500 characters) ---")
                    print(text[:500])
                    print("--------------------------------------")
                except Exception as e:
                    print(f"Failed to extract {file_path}: {e}")
